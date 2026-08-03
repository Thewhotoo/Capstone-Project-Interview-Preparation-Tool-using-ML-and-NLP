import docx
import pytest
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn

from resume_engine.extractor import ExtractionFailure, PdfDocxExtractor


def _add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    run_element = paragraph.add_run(text)._r
    paragraph._p.remove(run_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


@pytest.fixture
def plain_docx(tmp_path):
    path = tmp_path / "resume.docx"
    document = docx.Document()
    document.add_heading("Jordan Example", level=1)
    document.add_paragraph(
        "Software Engineer with several years of experience building distributed systems."
    )
    bold_paragraph = document.add_paragraph()
    run = bold_paragraph.add_run("Key Skills")
    run.bold = True
    document.save(str(path))
    return path


def test_extract_docx_returns_spans_with_text_and_style_name(plain_docx):
    result = PdfDocxExtractor().extract(str(plain_docx), "docx")
    texts = [s.text for s in result.spans]
    assert "Jordan Example" in texts
    assert any("Software Engineer" in t for t in texts)

    heading_span = next(s for s in result.spans if s.text == "Jordan Example")
    assert heading_span.style_name == "Heading 1"

    body_span = next(s for s in result.spans if "Software Engineer" in s.text)
    assert body_span.style_name == "Normal"


def test_extract_docx_detects_bold_runs(plain_docx):
    result = PdfDocxExtractor().extract(str(plain_docx), "docx")
    bold_span = next(s for s in result.spans if s.text == "Key Skills")
    assert bold_span.is_bold is True

    non_bold_span = next(s for s in result.spans if "Software Engineer" in s.text)
    assert non_bold_span.is_bold is False


def test_extract_docx_sets_source_format_and_page_count(plain_docx):
    result = PdfDocxExtractor().extract(str(plain_docx), "docx")
    assert result.source_format == "docx"
    assert result.page_count == 1


def test_extract_docx_computes_extraction_quality(plain_docx):
    result = PdfDocxExtractor().extract(str(plain_docx), "docx")
    assert result.extraction_quality.span_count == 3
    assert result.extraction_quality.chars_extracted > 0


def test_extract_docx_captures_hyperlinks(tmp_path):
    path = tmp_path / "resume_with_link.docx"
    document = docx.Document()
    document.add_paragraph("Jordan Example - a software engineer with hands-on production experience.")
    link_paragraph = document.add_paragraph()
    _add_hyperlink(link_paragraph, "https://linkedin.com/in/jordan", "LinkedIn Profile")
    document.save(str(path))

    result = PdfDocxExtractor().extract(str(path), "docx")

    assert len(result.hyperlinks) == 1
    url, bbox, page_num = result.hyperlinks[0]
    assert url == "https://linkedin.com/in/jordan"
    assert page_num == 0
    assert len(bbox) == 4
    assert any(s.text == "LinkedIn Profile" for s in result.spans)


def test_extract_docx_table_sidebar_assigns_distinct_column_geometry(tmp_path):
    path = tmp_path / "table_sidebar.docx"
    document = docx.Document()
    document.add_heading("Jordan Example", level=1)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).paragraphs[0].add_run(
        "Skills: Python, SQL, distributed systems design and production debugging."
    )
    table.cell(0, 1).paragraphs[0].add_run(
        "Experience: Acme Corp, Senior Engineer, led migration of the core billing platform."
    )
    document.save(str(path))

    result = PdfDocxExtractor().extract(str(path), "docx")

    skills_span = next(s for s in result.spans if "Skills:" in s.text)
    experience_span = next(s for s in result.spans if "Experience:" in s.text)
    assert skills_span.bbox[0] < experience_span.bbox[0]
    assert "docx_table_layout_detected:2_columns" in result.extraction_quality.notes


def test_extract_docx_raises_extraction_failure_for_corrupt_file(tmp_path):
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"this is not a real docx file")

    with pytest.raises(ExtractionFailure):
        PdfDocxExtractor().extract(str(path), "docx")


def test_extract_docx_raises_extraction_failure_for_near_empty_document(tmp_path):
    path = tmp_path / "sparse.docx"
    document = docx.Document()
    document.add_paragraph("Hi")
    document.save(str(path))

    with pytest.raises(ExtractionFailure):
        PdfDocxExtractor().extract(str(path), "docx")
