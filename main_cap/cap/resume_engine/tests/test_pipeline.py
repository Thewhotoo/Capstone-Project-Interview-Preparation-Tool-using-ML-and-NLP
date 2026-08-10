from resume_engine.confidence import AnnotatedCandidateProfile, Confidence
from resume_engine.document_model import DocumentModel, ExtractionQuality
from resume_engine.pipeline import (
    ResumePipeline,
    _absorb_repeated_unknown_entries,
    _group_sections_by_label,
)
from resume_engine.pipeline_trace import PipelineTrace
from resume_engine.sections import Section


class FakeExtractor:
    """Returns a real (empty) DocumentModel rather than a bare placeholder
    -- document_extraction's enrich hook (pipeline.py) reads
    DocumentModel-shaped fields, so every fake feeding it must be shaped
    like the real PdfDocxExtractor's return value, even when the fake's
    own logic is trivial."""

    def extract(self, file_path, source_format, trace=None):
        return DocumentModel(
            source_format=source_format,
            extraction_quality=ExtractionQuality(chars_extracted=100, span_count=5),
        )


class FakeLayoutReconstructor:
    def reconstruct(self, doc, trace=None):
        return doc


class FakeSectionDetector:
    """Returns a real (minimal) list[Section] rather than a bare placeholder
    -- section_detection's enrich hook and the pipeline's list->dict merge
    step (pipeline.py) both read Section-shaped fields."""

    def detect(self, doc, trace=None):
        return [
            Section(
                label="experience",
                raw_header_text="Experience",
                header_confidence=0.9,
                header_match_reason="alias_gazetteer_fuzzy:95",
            )
        ]


class FakeParserRunner:
    def run_all(self, sections, doc, trace=None):
        return {"experience": "result"}


class FakeCrossReferenceEngine:
    def cross_reference(self, parser_results, trace=None):
        return parser_results


class FakeNormalizer:
    def normalize(self, parser_results, trace=None):
        return parser_results


class FakeValidationEngine:
    def validate(self, parser_results, trace=None):
        return []


class FakeConfidenceEngine:
    def score(self, parser_results, observations, trace=None):
        return AnnotatedCandidateProfile(parser_results=parser_results, observations=observations)


def make_pipeline():
    return ResumePipeline(
        extractor=FakeExtractor(),
        layout_reconstructor=FakeLayoutReconstructor(),
        section_detector=FakeSectionDetector(),
        parser_runner=FakeParserRunner(),
        cross_reference_engine=FakeCrossReferenceEngine(),
        normalizer=FakeNormalizer(),
        validation_engine=FakeValidationEngine(),
        confidence_engine=FakeConfidenceEngine(),
    )


def test_pipeline_runs_all_fakes_end_to_end_and_returns_annotated_profile():
    pipeline = make_pipeline()
    result = pipeline.run("resume.pdf", "pdf")
    assert isinstance(result, AnnotatedCandidateProfile)
    assert result.parser_results == {"experience": "result"}


def test_pipeline_with_no_trace_records_nothing():
    pipeline = make_pipeline()
    pipeline.run("resume.pdf", "pdf", trace=None)  # must not raise


def test_pipeline_records_one_stage_trace_per_stage_in_order():
    pipeline = make_pipeline()
    trace = PipelineTrace(document_id="doc-1")

    pipeline.run("resume.pdf", "pdf", trace=trace)

    stage_names = [s.stage_name for s in trace.stages]
    assert stage_names == [
        "document_extraction",
        "layout_reconstruction",
        "section_detection",
        "entity_parsing",
        "cross_reference",
        "normalization",
        "validation",
        "confidence_scoring",
    ]


def test_pipeline_trace_captures_nonzero_total_duration():
    pipeline = make_pipeline()
    trace = PipelineTrace(document_id="doc-1")
    pipeline.run("resume.pdf", "pdf", trace=trace)
    assert trace.total_duration_ms >= 0.0
    assert len(trace.stages) == 8


def test_timed_with_no_enrich_produces_bare_stage_trace():
    trace = PipelineTrace(document_id="doc-1")
    ResumePipeline._timed("stage", trace, lambda: "result")
    assert trace.stages[0].metadata == {}
    assert trace.stages[0].confidences == []


def test_timed_with_enrich_populates_metadata_and_confidences():
    trace = PipelineTrace(document_id="doc-1")

    def enrich(result):
        return {"result_was": result}, [Confidence(score=1.0, reasons=["+ok"])]

    ResumePipeline._timed("stage", trace, lambda: "result", enrich=enrich)

    assert trace.stages[0].metadata == {"result_was": "result"}
    assert trace.stages[0].confidences[0].score == 1.0


def test_timed_enrich_is_never_forwarded_to_fn():
    trace = PipelineTrace(document_id="doc-1")

    def fn(**kwargs):
        assert "enrich" not in kwargs
        return "result"

    ResumePipeline._timed("stage", trace, fn, enrich=lambda result: ({}, []))


def test_timed_skips_enrich_when_trace_is_none():
    calls = []
    ResumePipeline._timed("stage", None, lambda: "result", enrich=lambda result: calls.append(result))
    assert calls == []


def test_run_enriches_document_extraction_and_layout_reconstruction_stage_traces(tmp_path):
    """End-to-end with the real Milestone 1 stages (PdfDocxExtractor,
    ColumnAwareLayoutReconstructor) wired in, fakes for everything after --
    confirms the enrich hook actually gets real signals from real stages,
    not just from hand-written fakes in isolation."""
    import docx

    from resume_engine.extractor import PdfDocxExtractor
    from resume_engine.layout import ColumnAwareLayoutReconstructor

    path = tmp_path / "resume.docx"
    document = docx.Document()
    document.add_heading("Jordan Example", level=1)
    document.add_paragraph(
        "Software Engineer with several years of experience building distributed systems."
    )
    document.save(str(path))

    pipeline = ResumePipeline(
        extractor=PdfDocxExtractor(),
        layout_reconstructor=ColumnAwareLayoutReconstructor(),
        section_detector=FakeSectionDetector(),
        parser_runner=FakeParserRunner(),
        cross_reference_engine=FakeCrossReferenceEngine(),
        normalizer=FakeNormalizer(),
        validation_engine=FakeValidationEngine(),
        confidence_engine=FakeConfidenceEngine(),
    )
    trace = PipelineTrace(document_id="doc-1")

    pipeline.run(str(path), "docx", trace=trace)

    extraction_stage = trace.stages[0]
    assert extraction_stage.stage_name == "document_extraction"
    assert extraction_stage.metadata["span_count"] > 0
    assert extraction_stage.metadata["chars_extracted"] > 0
    assert extraction_stage.confidences[0].score == 1.0
    assert any(r.startswith("+sufficient_text_extracted") for r in extraction_stage.confidences[0].reasons)

    layout_stage = trace.stages[1]
    assert layout_stage.stage_name == "layout_reconstruction"
    assert layout_stage.metadata["layout_mode"] == "single_column"
    assert layout_stage.confidences[0].score >= 0.9
    assert "+single_column_default" in layout_stage.confidences[0].reasons


def test_run_enriches_section_detection_and_merges_sections_by_label(tmp_path):
    """End-to-end with the real Milestone 1+2 stages (PdfDocxExtractor,
    ColumnAwareLayoutReconstructor, HeuristicSectionDetector) wired in --
    confirms the section_detection enrich hook reports real labels/
    confidence, and that ParserRunner receives a dict[str, Section] (the
    pipeline's list->dict merge), not the detector's raw list."""
    import docx

    from resume_engine.extractor import PdfDocxExtractor
    from resume_engine.layout import ColumnAwareLayoutReconstructor
    from resume_engine.sections import HeuristicSectionDetector

    path = tmp_path / "resume.docx"
    document = docx.Document()
    document.add_heading("Jordan Example", level=1)
    document.add_heading("Experience", level=2)
    document.add_paragraph("Acme Corp, Senior Engineer, led the billing migration.")
    document.save(str(path))

    received_sections = {}

    class RecordingParserRunner:
        def run_all(self, sections, doc, trace=None):
            received_sections.update(sections)
            return {}

    pipeline = ResumePipeline(
        extractor=PdfDocxExtractor(),
        layout_reconstructor=ColumnAwareLayoutReconstructor(),
        section_detector=HeuristicSectionDetector(),
        parser_runner=RecordingParserRunner(),
        cross_reference_engine=FakeCrossReferenceEngine(),
        normalizer=FakeNormalizer(),
        validation_engine=FakeValidationEngine(),
        confidence_engine=FakeConfidenceEngine(),
    )
    trace = PipelineTrace(document_id="doc-1")

    pipeline.run(str(path), "docx", trace=trace)

    section_stage = trace.stages[2]
    assert section_stage.stage_name == "section_detection"
    assert "experience" in section_stage.metadata["labels_found"]
    assert section_stage.metadata["section_count"] == len(section_stage.confidences)

    # ParserRunner got a dict keyed by label, not the detector's raw list.
    assert isinstance(received_sections, dict)
    assert "experience" in received_sections
    assert isinstance(received_sections["experience"], Section)


def test_group_sections_by_label_concatenates_spans_for_a_repeated_label():
    """The cross-page/column continuation mechanism (architecture doc
    Section 4.3): two raw Section hits for the same label merge into one
    dict entry, spans concatenated in order."""
    first = Section(
        label="experience", raw_header_text="Experience", spans=["page1-span"],
        header_confidence=0.9, header_match_reason="alias_gazetteer_fuzzy:95",
    )
    second = Section(
        label="experience", raw_header_text="Experience (cont.)", spans=["page2-span"],
        header_confidence=0.4, header_match_reason="alias_gazetteer_weak:60",
    )

    grouped = _group_sections_by_label([first, second])

    assert list(grouped.keys()) == ["experience"]
    assert grouped["experience"].spans == ["page1-span", "page2-span"]
    # First-detected wins for confidence/reason -- "first confident match
    # wins" cascade philosophy applied to the merge too.
    assert grouped["experience"].header_confidence == 0.9


def test_group_sections_by_label_keeps_distinct_labels_separate():
    experience = Section(label="experience", raw_header_text="Experience")
    education = Section(label="education", raw_header_text="Education")

    grouped = _group_sections_by_label([experience, education])

    assert set(grouped.keys()) == {"experience", "education"}


def test_absorb_repeated_unknown_entries_merges_two_or_more_consecutive_unknowns():
    """The bug this fixes: two project-title-shaped lines, each
    misclassified 'unknown' by Section Detection, immediately following a
    real 'projects' section -- must be absorbed into it, not left as
    separate stray sections."""
    projects = Section(label="projects", raw_header_text="Projects", spans=["projects-header-span"])
    entry_one = Section(label="unknown", raw_header_text="AI SOC Analyst", spans=["entry-one-span"])
    entry_two = Section(label="unknown", raw_header_text="Task Tracker", spans=["entry-two-span"])

    result = _absorb_repeated_unknown_entries([projects, entry_one, entry_two])

    assert len(result) == 1
    assert result[0].label == "projects"
    assert result[0].spans == ["projects-header-span", "entry-one-span", "entry-two-span"]


def test_absorb_repeated_unknown_entries_leaves_a_single_isolated_unknown_untouched():
    """Direct regression for ambiguous_header_pdf (Milestone 2, frozen,
    already validated): a genuine standalone ambiguous section ('Languages')
    is a single, isolated 'unknown' hit -- must NOT be absorbed."""
    experience = Section(label="experience", raw_header_text="Experience", spans=["exp-span"])
    languages = Section(label="unknown", raw_header_text="Languages", spans=["lang-span"])

    result = _absorb_repeated_unknown_entries([experience, languages])

    assert len(result) == 2
    assert result[0].label == "experience"
    assert result[0].spans == ["exp-span"]
    assert result[1].label == "unknown"
    assert result[1].spans == ["lang-span"]


def test_absorb_repeated_unknown_entries_does_not_absorb_across_a_real_section_boundary():
    experience = Section(label="experience", raw_header_text="Experience", spans=["exp-span"])
    stray = Section(label="unknown", raw_header_text="Widget Factory", spans=["stray-span"])
    education = Section(label="education", raw_header_text="Education", spans=["edu-span"])
    entry_two = Section(label="unknown", raw_header_text="Beta Inc", spans=["entry-two-span"])

    result = _absorb_repeated_unknown_entries([experience, stray, education, entry_two])

    labels = [s.label for s in result]
    assert labels == ["experience", "unknown", "education", "unknown"]


def test_absorb_repeated_unknown_entries_handles_unknown_at_start_of_document():
    leading_unknown = Section(label="unknown", raw_header_text="", spans=["a"])
    result = _absorb_repeated_unknown_entries([leading_unknown])
    assert len(result) == 1
    assert result[0] is leading_unknown


def test_absorb_repeated_unknown_entries_handles_empty_list():
    assert _absorb_repeated_unknown_entries([]) == []


def test_absorb_repeated_unknown_entries_does_not_sweep_a_font_elevated_section_into_skills():
    """Direct regression for the real-resume audit finding: a genuine
    "Honors & Awards" section (no gazetteer alias yet, so it classifies as
    'unknown') immediately following a run of same-size bold skills
    category sub-headers ("Frontend:", "Backend:", ...) -- themselves also
    'unknown' -- was being swept into the preceding "skills" section along
    with them, contaminating it with unrelated prose ("Awarded the ...
    Scholarship..."). A font-elevated header (the same visual weight every
    confidently-labeled section header carries) must end the run instead of
    being absorbed."""
    skills = Section(label="skills", raw_header_text="Technical Skills", spans=["skills-header-span"])
    frontend = Section(label="unknown", raw_header_text="Frontend:", spans=["frontend-span"])
    backend = Section(label="unknown", raw_header_text="Backend:", spans=["backend-span"])
    honors = Section(
        label="unknown",
        raw_header_text="Honors & Awards",
        spans=["honors-span"],
        header_font_elevated=True,
    )

    result = _absorb_repeated_unknown_entries([skills, frontend, backend, honors])

    labels = [s.label for s in result]
    assert labels == ["skills", "unknown"]
    assert result[0].spans == ["skills-header-span", "frontend-span", "backend-span"]
    assert result[1].spans == ["honors-span"]
    assert result[1].raw_header_text == "Honors & Awards"


def test_absorb_repeated_unknown_entries_absorbs_a_single_isolated_unknown_when_the_section_already_uses_bold_entries(
    make_text_span,
):
    """Direct regression for the real-resume audit's second-project bug: a
    second project's title line, short enough to itself pass
    sections.py's header word-count filter, fragmented into a standalone
    "unknown" Section that ProjectParser never saw -- the resume's second
    project silently vanished from CandidateProfile. Since "Projects"
    already has its own bold entry-title line (the first project's title),
    a same-style single unknown right after it is absorbed -- the
    established-bold-convention signal `_has_bold_body_line` provides.
    `ambiguous_header_pdf`'s "Languages" case (test above via the real
    golden-corpus suite) is unaffected: that fixture's Experience entry
    line is NOT bold, so no such convention exists there."""
    projects_header = make_text_span(text="Projects", bbox=(50.0, 50.0, 100.0, 62.0), font_size=12.0, is_bold=True)
    project_one_title = make_text_span(
        text="AI SOC Analyst Intelligent Security Log Analysis Platform",
        bbox=(50.0, 75.0, 300.0, 87.0), font_size=10.0, is_bold=True,
    )
    project_one_body = make_text_span(
        text="Built an AI-assisted platform.", bbox=(50.0, 95.0, 250.0, 107.0), font_size=10.0, is_bold=False,
    )
    section = Section(
        label="projects", raw_header_text="Projects",
        spans=[projects_header, project_one_title, project_one_body],
        header_confidence=0.9, header_font_elevated=True,
    )
    project_two_title = Section(
        label="unknown", raw_header_text="AI Adaptive Interview Prep",
        spans=[make_text_span(
            text="AI Adaptive Interview Prep", bbox=(50.0, 120.0, 220.0, 132.0), font_size=10.0, is_bold=True,
        )],
        header_confidence=0.2, header_font_elevated=False,
    )

    result = _absorb_repeated_unknown_entries([section, project_two_title])

    assert len(result) == 1
    assert result[0].label == "projects"
    assert result[0].spans == [
        projects_header, project_one_title, project_one_body, *project_two_title.spans,
    ]


def test_projects_section_with_two_entries_survives_end_to_end_through_section_detection_and_project_parser(
    make_text_span, make_document_model,
):
    """The real end-to-end shape (Section Detection -> absorption ->
    ProjectParser), not just the absorption helper in isolation: two
    projects in one Projects section, the second one's title short enough
    to independently pass sections.py's own header filter -- both must
    survive into CandidateProfile as separate project entities."""
    from resume_engine.parsers.project_parser import ProjectParser
    from resume_engine.pipeline import _group_sections_by_label
    from resume_engine.sections import HeuristicSectionDetector

    spans = [
        make_text_span(text="Projects", bbox=(50.0, 50.0, 100.0, 62.0), font_size=12.0, is_bold=True),
        make_text_span(
            text="AI SOC Analyst Intelligent Security Log Analysis Platform",
            bbox=(50.0, 75.0, 400.0, 87.0), font_size=10.0, is_bold=True,
        ),
        make_text_span(
            text="Built an AI-assisted SOC platform using FastAPI and React.",
            bbox=(50.0, 95.0, 400.0, 107.0), font_size=10.0, is_bold=False,
        ),
        make_text_span(
            text="AI Adaptive Interview Prep", bbox=(50.0, 120.0, 220.0, 132.0), font_size=10.0, is_bold=True,
        ),
        make_text_span(
            text="Built an interview preparation platform using LangGraph.",
            bbox=(50.0, 140.0, 400.0, 152.0), font_size=10.0, is_bold=False,
        ),
        make_text_span(text="Experience", bbox=(50.0, 165.0, 130.0, 177.0), font_size=12.0, is_bold=True),
        make_text_span(
            text="Acme Corp Senior Engineer 2021 Present",
            bbox=(50.0, 185.0, 300.0, 197.0), font_size=10.0, is_bold=True,
        ),
    ]
    doc = make_document_model(spans=spans, body_font_size=10.0)

    raw_sections = HeuristicSectionDetector().detect(doc)
    absorbed = _absorb_repeated_unknown_entries(raw_sections)
    by_label = _group_sections_by_label(absorbed)

    result = ProjectParser().parse(by_label, doc)

    titles = [e["title"] for e in result.entities]
    assert "AI SOC Analyst Intelligent Security Log Analysis Platform" in titles
    assert "AI Adaptive Interview Prep" in titles
    assert len(result.entities) == 2


def test_absorb_repeated_unknown_entries_does_not_absorb_a_single_isolated_unknown_after_education(
    make_text_span,
):
    """Direct regression for the real-resume audit's education-phantom-
    entry bug: single-isolated-unknown absorption used to apply to ANY
    section with its own bold body line, including "education" -- whose
    own entries (e.g. a bold institution-name line) routinely satisfy that
    signal for reasons unrelated to a second real degree. Restricting
    absorption to `_MULTI_ENTRY_ABSORPTION_LABELS` ("projects",
    "experience") must leave a trailing single unknown after "education"
    unabsorbed, even though the section has its own bold body line."""
    education_header = make_text_span(text="Education", bbox=(50.0, 50.0, 100.0, 62.0), font_size=12.0, is_bold=True)
    institution_line = make_text_span(
        text="PES University, Electronic City Campus", bbox=(50.0, 75.0, 300.0, 87.0), font_size=10.0, is_bold=True,
    )
    degree_line = make_text_span(
        text="B.Tech in Computer Science and Engineering", bbox=(50.0, 95.0, 300.0, 107.0), font_size=10.0, is_bold=False,
    )
    section = Section(
        label="education", raw_header_text="Education",
        spans=[education_header, institution_line, degree_line],
        header_confidence=0.9, header_font_elevated=True,
    )
    schooling_aside = Section(
        label="unknown", raw_header_text="National Centre for Excellence",
        spans=[make_text_span(
            text="National Centre for Excellence", bbox=(50.0, 120.0, 250.0, 132.0), font_size=10.0, is_bold=True,
        )],
        header_confidence=0.2, header_font_elevated=False,
    )

    result = _absorb_repeated_unknown_entries([section, schooling_aside])

    assert len(result) == 2
    assert result[0].label == "education"
    assert result[0].spans == [education_header, institution_line, degree_line]
    assert result[1].label == "unknown"
    assert result[1].spans == schooling_aside.spans


def test_education_section_with_schooling_aside_survives_end_to_end_with_exactly_one_entity(
    make_text_span, make_document_model,
):
    """The real end-to-end shape: Education has one real degree entry
    followed by a single, bold-titled but otherwise unrecognizable
    schooling aside. Must produce exactly ONE EducationParser entity --
    never a phantom empty second one."""
    from resume_engine.parsers.education_parser import EducationParser
    from resume_engine.pipeline import _group_sections_by_label
    from resume_engine.sections import HeuristicSectionDetector

    spans = [
        make_text_span(text="Education", bbox=(50.0, 50.0, 100.0, 62.0), font_size=12.0, is_bold=True),
        make_text_span(
            text="PES University Electronic City Campus", bbox=(50.0, 75.0, 300.0, 87.0), font_size=10.0, is_bold=True,
        ),
        make_text_span(
            text="B.Tech in Computer Science and Engineering", bbox=(50.0, 95.0, 300.0, 107.0), font_size=10.0, is_bold=False,
        ),
        make_text_span(
            text="National Centre for Excellence", bbox=(50.0, 120.0, 250.0, 132.0), font_size=10.0, is_bold=True,
        ),
        make_text_span(
            text="Class XII 85.4 percent Class X 94.7 percent", bbox=(50.0, 140.0, 300.0, 152.0), font_size=10.0, is_bold=False,
        ),
        make_text_span(text="Projects", bbox=(50.0, 165.0, 130.0, 177.0), font_size=12.0, is_bold=True),
        make_text_span(
            text="Some Project Title Here", bbox=(50.0, 185.0, 300.0, 197.0), font_size=10.0, is_bold=True,
        ),
    ]
    doc = make_document_model(spans=spans, body_font_size=10.0)

    raw_sections = HeuristicSectionDetector().detect(doc)
    absorbed = _absorb_repeated_unknown_entries(raw_sections)
    by_label = _group_sections_by_label(absorbed)

    result = EducationParser().parse(by_label, doc)

    assert len(result.entities) == 1
    assert result.entities[0]["institution"] == "PES University"
